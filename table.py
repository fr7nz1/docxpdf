import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK

# doc = docx.Document('1.docx')
doc = docx.Document('Тест.docx')
# doc = docx.Document('Тест(табл).docx')


def check_num_tables(file):
    tablescount = 0
    # Проходим по всем таблицам в документе
    for table in doc.tables[1:]:
        tablescount += 1
        pass
    # print("Кол-во таблиц:", tablescount)
    return tablescount


def check_link_tables(file):
    linktablescount = 0
    i = 1
    mask_template = """Таблица {chislo} –"""  # Задаем маску для последующей проверки по ней
    mask = mask_template.format(chislo=i)  # Добавляет переменную в маску
    for paragraph in doc.paragraphs:  # Проходит по всем абзацам
        if mask in paragraph.text:  # Проверяет есть ли маска в тексте
            linktablescount += 1
            i += 1
            mask = mask_template.format(chislo=i)
            # if linktablescount >= tablescount:
                # тут мы делаем пометку в данном параграфе что ссылок стало больше чем таблиц
                # comment = paragraph.add_comment("Кол-во ссылок на таблицы и кол-во таблиц не совпадает!")
                # comment.author = 'bot'
    # print("Кол-во ссылок на таблицы:", linktablescount)
    return linktablescount


def check_num_pic(file):
    piccount = 0
    # Проходит по всем иллюстрациям в документе
    for shape in doc.inline_shapes:
        piccount += 1
        pass
    # print("Кол-во иллюстраций:", piccount)
    return piccount


def check_link_pic(file):
    linkpiccount = 0
    i = 1
    mask_template_pic = """Рисунок {chislo} –"""  # Задаем маску для последующей проверки по ней
    mask_pic = mask_template_pic.format(chislo=i)  # Добавляет переменную в маску
    for paragraph in doc.paragraphs:  # Проходит по всем абзацам
        if mask_pic in paragraph.text:  # Проверяет есть ли маска в тексте
            linkpiccount += 1
            i += 1
            mask_pic = mask_template_pic.format(chislo=i)
    # print("Кол-во ссылок на иллюстрации:", linkpiccount)
    return linkpiccount


def check_num_sources(file):
    sourcescount = 0

    flag = False
    mask_title = """Список литературы"""

    for paragraph in doc.paragraphs:
        if mask_title in paragraph.text:
            flag = True

        if flag == True:
            if paragraph.style.name == 'List Paragraph':
                sourcescount += 1
                print(paragraph.text)
    # i = 1
    # mask_template_sources = """{chislo}. """
    # mask_sources = mask_template_sources.format(chislo=i)
    #
    # for paragraph in doc.paragraphs:  # Нужно придумать как начать сразу с последней страницы
    #     if mask_sources in paragraph.text:
    #         sourcescount += 1
    #         i += 1
    #         mask_sources = mask_template_sources.format(chislo=i)
    # print("Кол-во источников:", sourcescount)
    return sourcescount


def check_link_sources(file):
    linksourcecount = 0
    i = 1
    mask_template_link_sources = """[{chislo}]"""
    mask_link_sources = mask_template_link_sources.format(chislo=i)

    for paragraph in doc.paragraphs:  # Нужно придумать как начать сразу с последней страницы
        if mask_link_sources in paragraph.text:  # Проверяет есть ли маска в тексте
            linksourcecount += 1
            i += 1
            mask_link_sources = mask_template_link_sources.format(chislo=i)
    # print("Кол-во ссылок на источники:", linksourcecount)
    return linksourcecount


def check_sources_merged(file):
    sourcescount = 0
    linksourcecount = 0

    flag = False
    mask_title = """Список литературы"""  # По госту должен быть БИБЛИОГРАФИЧЕСКИЙ СПИСОК
    # mask_title = """БИБЛИОГРАФИЧЕСКИЙ СПИСОК"""

    j = 1
    mask_template_link_sources = """[{chislo}]"""
    mask_link_sources = mask_template_link_sources.format(chislo=j)

    for paragraph in doc.paragraphs:    # Доходим до заголовка Список литературы
        if mask_title in paragraph.text:
            flag = True
        if flag == True:    # Считаем количество источников
            if paragraph.style.name == 'List Paragraph':
                sourcescount += 1

    spisok = []
    for i in range(1, sourcescount + 1):
        spisok.append(0)

    for paragraph in doc.paragraphs:
        j = 1
        mask_link_sources = mask_template_link_sources.format(chislo=j)
        for i in range(1, sourcescount + 1):
            if mask_link_sources in paragraph.text:
                linksourcecount += 1
                spisok[i - 1] = i
                j += 1
                mask_link_sources = mask_template_link_sources.format(chislo=j)
            else:
                j += 1
                mask_link_sources = mask_template_link_sources.format(chislo=j)

    for paragraph in doc.paragraphs:
        s = 0
        for i in spisok:
            s += 1
            if i == 0:
                comment = paragraph.add_comment("Отсутствует ссылка на источник:", str(s))
                comment.author = 'bot'
                print("Отсутствует ссылка на источник:", s)
        break
    # print(sourcescount, linksourcecount)
    # print(spisok)


def check_pic_merged(file):
    piccount = 0
    linkpiccount = 0
    j = 1
    mask_template_pic = """Рисунок {chislo} –"""  # Задаем маску для последующей проверки по ней
    mask_pic = mask_template_pic.format(chislo=j)  # Добавляет переменную в маску

    for shape in doc.inline_shapes:     # Проходит по всем иллюстрациям в документе
        piccount += 1
        pass

    spisok = []
    for i in range(1, piccount + 1):
        spisok.append(0)

    for paragraph in doc.paragraphs:
        j = 1
        mask_pic = mask_template_pic.format(chislo=j)
        for i in range(1, piccount + 1):
            if mask_pic in paragraph.text:
                linkpiccount += 1
                spisok[i - 1] = i
                j += 1
                mask_pic = mask_template_pic.format(chislo=j)
            else:
                j += 1
                mask_pic = mask_template_pic.format(chislo=j)

    for paragraph in doc.paragraphs:
        s = 0
        for i in spisok:
            s += 1
            if i == 0:
                comment = paragraph.add_comment("Отсутствует ссылка на иллюстрацию:", str(s))
                comment.author = 'bot'
                print("Отсутствует ссылка на иллюстрацию:", s)
        break
    # print(piccount, linkpiccount)
    # print(spisok)


def check_table_merged(file):
    tablescount = 0
    linktablescount = 0
    j = 1
    mask_template = """Таблица {chislo} –"""  # Задаем маску для последующей проверки по ней
    mask = mask_template.format(chislo=j)  # Добавляет переменную в маску

    for table in doc.tables:  # for table in doc.tables[1:]:    # Проходим по всем таблицам в документе
        tablescount += 1
        pass

    spisok = []
    for i in range(1, tablescount + 1):
        spisok.append(0)

    for paragraph in doc.paragraphs:
        j = 1
        mask = mask_template.format(chislo=j)
        for i in range(1, tablescount + 1):
            if mask in paragraph.text:
                linktablescount += 1
                spisok[i - 1] = i
                j += 1
                mask = mask_template.format(chislo=j)
            else:
                j += 1
                mask = mask_template.format(chislo=j)

    for paragraph in doc.paragraphs:
        s = 0
        for i in spisok:
            s += 1
            if i == 0:
                comment = paragraph.add_comment("Отсутствует ссылка на таблицу:", str(s))
                comment.author = 'bot'
                print("Отсутствует ссылка на таблицу:", s)
        break
    # print(tablescount, linktablescount)
    # print(spisok)


def bot_comment(paragraph, TEXT: str):
    comment = paragraph.add_comment(TEXT)
    comment.author = 'bot'


def check_block(file):
    referat = False
    mask_referat = """РЕФЕРАТ"""

    soderjanie = False
    mask_soderjanie = """СОДЕРЖАНИЕ"""

    vvedenie = False
    mask_vvedenie = """ВВЕДЕНИЕ"""

    zakluchenie = False
    mask_zakluchenie = """ЗАКЛЮЧЕНИЕ"""

    bibliogr_spisok = False
    mask_bibliogr_spisok = """БИБЛИОГРАФИЧЕСКИЙ СПИСОК"""

    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1':
            if mask_referat in paragraph.text:
                referat = True
                # print(paragraph.text)
            elif mask_soderjanie in paragraph.text:
                soderjanie = True
                # print(paragraph.text)
            elif mask_vvedenie in paragraph.text:
                vvedenie = True
                # print(paragraph.text)
            elif mask_zakluchenie in paragraph.text:
                zakluchenie = True
                # print(paragraph.text)
            elif mask_bibliogr_spisok in paragraph.text:
                bibliogr_spisok = True
                # print(paragraph.text)

    # if (referat != True):
    #     for paragraph in doc.paragraphs:
    #         comment = paragraph.add_comment('Блок РЕФЕРАТ отсутствует!')
    #         comment.author = 'bot'
    #         break
    #     print("Блок РЕФЕРАТ отсутствует!")

    if (referat != True):
        for paragraph in doc.paragraphs:
            bot_comment(paragraph, "Блок РЕФЕРАТ отсутствует!")
            break
        print("Блок РЕФЕРАТ отсутствует!")

    if (soderjanie != True):
        for paragraph in doc.paragraphs:
            bot_comment(paragraph, "Блок СОДЕРЖАНИЕ отсутствует!")
            break
        print("Блок СОДЕРЖАНИЕ отсутствует!")

    if (vvedenie != True):
        for paragraph in doc.paragraphs:
            bot_comment(paragraph, "Блок ВВЕДЕНИЕ отсутствует!")
            break
        print("Блок ВВЕДЕНИЕ отсутствует!")

    if (zakluchenie != True):
        for paragraph in doc.paragraphs:
            bot_comment(paragraph, "Блок ЗАКЛЮЧЕНИЕ отсутствует!")
            break
        print("Блок ЗАКЛЮЧЕНИЕ отсутствует!")

    if (bibliogr_spisok != True):
        for paragraph in doc.paragraphs:
            bot_comment(paragraph, "Блок БИБЛИОГРАФИЧЕСКИЙ СПИСОК отсутствует!")
            break
        print("Блок БИБЛИОГРАФИЧЕСКИЙ СПИСОК отсутствует!")


# check_sources_merged(doc)
# check_table_merged(doc)
# check_pic_merged(doc)

# check_num_sources(doc)
# check_link_sources(doc)

# check_num_tables(doc)
# check_link_tables(doc)

# check_num_pic(doc)
# check_link_pic(doc)

# check_block(doc)