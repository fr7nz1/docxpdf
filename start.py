import docx
import functions.margin
import functions.formatting
import functions.table

import os

paths = []


def search():
    # Ищет доки
    for file in os.listdir(os.getcwd()):
        if file.endswith('.docx'):
            paths.append(file)


def names(path):
    text = 'Наименование документа: ' + path
    return text



if __name__ == '__main__':
    # выполняет все функции
    try:
        # Ищет доки
        search()
        for path in paths:
            # отдает док в переменную
            doc = docx.Document(path)
            print(names(path))

            functions.formatting.indent_img(doc.paragraphs)  # написание в первой строке, о проверке отступау "Рисунок № -"

            functions.margin.margin(doc)  # проверка полей

            functions.formatting.check_size_font(177800, doc.paragraphs)  # проверка размера шрифта
            functions.formatting.check_name_font('Times New Roman', doc.paragraphs)  # проверка имени шрифта

            functions.formatting.check_heading(1.25, doc.paragraphs)  # проверка заголовков
            functions.formatting.check_first_line_indent(1.25, doc.paragraphs)  # проверка отступа первого абзаца
            functions.formatting.check_line_spacing(1.5, doc.paragraphs)  # проверка отступа межстрочного

            functions.table.check_sources_merged(doc)  # библ список?
            functions.table.check_pic_merged(doc)    # изображения
            functions.table.check_table_merged(doc)  # таблицы
            doc.save(path)
    except Exception as e:
        print('Ошибка!')
